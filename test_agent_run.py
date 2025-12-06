# test_agent_run.py
# This script creates example PDF and Word documents, then runs the AI Research Agent

import os
from pathlib import Path
from docx import Document
from PyPDF2 import PdfWriter
import asyncio
from main import main as run_agent

# Ensure docs folder exists
Path("rag/docs").mkdir(parents=True, exist_ok=True)

# --- Create example PDF ---
pdf_path = "rag/docs/doc1.pdf"
writer = PdfWriter()
writer.add_blank_page(width=72*8.5, height=72*11)  # 8.5x11 inches
with open(pdf_path, "wb") as f:
    writer.write(f)
# Since blank page, we will just simulate text by FAISS fallback

# --- Create example Word doc ---
docx_path = "rag/docs/doc2.docx"
doc = Document()
doc.add_paragraph("AI applications in NLP include chatbots for patient engagement and automated medical transcription.")
doc.add_paragraph("Deep learning enables prediction of patient outcomes and risk assessment.")
doc.add_paragraph("Healthcare providers benefit from AI-driven decision support systems.")
doc.save(docx_path)

# --- Run the agent ---
asyncio.run(run_agent())
